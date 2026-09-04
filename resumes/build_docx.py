"""
Generate four role-specific DOCX resumes for editing.
Run: python build_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def make_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(12)
    return doc


def add_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MARTY SCOTT")
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(they/them)")
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("Brooklyn, NY  |  909.456.0467  |  greatmartyscott@gmail.com")
    run.font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("marty64.net  |  github.com/martychouette  |  martychouette.itch.io")
    run.font.size = Pt(9.5)


def add_summary(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(9.5)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Thin separator line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    run2 = p2.add_run("_" * 95)
    run2.font.size = Pt(2)
    run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def add_role(doc, title, detail=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    if detail:
        run2 = p.add_run(f"  |  {detail}")
        run2.font.size = Pt(9.5)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    for run in p.runs:
        run.clear()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"


def add_skills_line(doc, label, detail):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(9.5)
    run2 = p.add_run(detail)
    run2.font.size = Pt(9.5)


def add_edu(doc, school, degree):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(f"{school} - ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(degree)
    run2.font.size = Pt(10)


EDUCATION = [
    ("NYU (Tisch)", "MFA, Game Design (Game Center) | 2024 - 2026"),
    ("Cal Poly Pomona", "BS, Computer Science | 2021 - 2023"),
    ("Mt. San Antonio College", "AA/AS, Database Management & Programming | 2018 - 2021"),
    ("Musicians Institute", "AA, Piano Tech (Audio Engineering emphasis) | 2009 - 2011"),
]


def add_education(doc):
    add_section_heading(doc, "Education")
    for school, degree in EDUCATION:
        add_edu(doc, school, degree)


# ============================================================
# 1. TEACHING
# ============================================================
def build_teaching():
    doc = make_doc()
    add_header(doc)
    add_summary(doc,
        "A computer scientist and game designer who treats code as creative material. "
        "I build interactive systems that explore what code can express - sightless platformers "
        "navigated by sound, autobiographical physics games, and an aesthetic-first game engine "
        "in C++/Vulkan. I develop open pedagogical frameworks for teaching code to non-programmers "
        "through making, with accessibility and craft at the center of everything I teach."
    )

    add_section_heading(doc, "Skills")
    add_skills_line(doc, "Pedagogy", "Curriculum design, scaffolded instruction, studio-based teaching, accessibility-centered design")
    add_skills_line(doc, "Languages", "C#, C++, Python, JavaScript, Java, SQL")
    add_skills_line(doc, "Engines / Tools", "Unity, Unreal, GameMaker, p5.js, Processing")
    add_skills_line(doc, "Game Design", "Rapid prototyping, play-testing, iterative design, system balancing")
    add_skills_line(doc, "Audio", "FMOD, Pro Tools, Logic, Ableton (15+ years)")

    add_section_heading(doc, "Teaching Experience")

    add_role(doc, "NYU Game Center (Tisch) - Code Help Desk", "2025 - Present")
    add_bullet(doc, "Help BFA and MFA students with code issues across Unity, Unreal, Processing, Twine, and GameMaker")
    add_bullet(doc, "Walk students through technical problems in plain language so they can get back to making")
    add_bullet(doc, "Document recurring fixes for the desk")

    add_role(doc, "NYU Game Center Conference - Speaker", "2025")
    add_bullet(doc, "Presented on game design pedagogy and accessibility")

    add_role(doc, "Music & Arts - Music Teacher", "2012 - 2018")
    add_bullet(doc, "Taught piano, guitar, and sound production across age groups and skill levels")
    add_bullet(doc, "Six years of one-on-one and small group instruction")

    add_section_heading(doc, "Pedagogical Work")

    add_role(doc, "VG101 Knowledge Base", "marty64.net/vg101 | In Development")
    add_bullet(doc, "Open pedagogical text for teaching videogame design as its own medium, built on a Gesture + Aesthetic Heritage framework")
    add_bullet(doc, 'Core pages: "Code as Material" (Schon, Papert, Sennett), "Debugging as Literacy," "Accessibility as Craft"')
    add_bullet(doc, "Teaching cycle: Play - Name - Make - Reflect, grounded in Dewey and Kolb")
    add_bullet(doc, "Three-register system (Practice / Craft / Theory) for layered reading across student and educator audiences")
    add_bullet(doc, "Code Bank: pre-shaped scaffolds students tune before they write, with exposed parameters for exploration")

    add_section_heading(doc, "Selected Projects")

    add_role(doc, "Blackout", "Solo Developer | Unity, FMOD")
    add_bullet(doc, "Sightless 2D platformer built around auditory navigation with no visual cues")
    add_bullet(doc, "Published research whitepaper on designing for blind play")
    add_bullet(doc, "Directly engages software and disability as a design question, not a compliance checkbox")

    add_role(doc, "Prototype Studio Gallery", "11 Weekly Prototypes | Unity")
    add_bullet(doc, "One prototype per week, each exploring a different question through code")
    add_bullet(doc, "Text adventures, physics toys, 3D environments, naval strategy, emotional vignettes")

    add_role(doc, "Cold Flame", "Solo Developer | Unity")
    add_bullet(doc, "Autobiographical physics game about life, loss, and chopping wood; code as personal expression")

    add_role(doc, "TEGE - The Enjin Game Engine", "Solo Developer | C++20, Vulkan")
    add_bullet(doc, "Aesthetic-first open source game engine with a complete editor, 70+ ECS components, and AngelScript scripting")
    add_bullet(doc, "Built a comprehensive accessibility suite: 8 colorblind modes, screen reader support, switch access, dyslexia mode")
    add_bullet(doc, "Licensed BSL 1.1 with full documentation and installer")

    add_education(doc)

    doc.save(os.path.join(OUT_DIR, "marty-scott-teaching.docx"))
    print("Built: marty-scott-teaching.docx")


# ============================================================
# 2. PROGRAMMING
# ============================================================
def build_programming():
    doc = make_doc()
    add_header(doc)
    add_summary(doc,
        "A software engineer and game designer with a CS degree and an MFA from NYU. "
        "I build game engines, gameplay systems, and tools from the ground up - from a "
        "100k+ line Vulkan engine with ray tracing and accessibility features, to narrative "
        "state machines and physics-driven interactions in Unity and Unreal. I care about "
        "code that feels good to use, looks right on screen, and ships."
    )

    add_section_heading(doc, "Skills")
    add_skills_line(doc, "Languages", "C#, C++, Java, Python, JavaScript, SQL, Custom Scripting")
    add_skills_line(doc, "Engines / Tools", "Unity, Unreal, GameMaker, TEGE (open source, BSL 1.1)")
    add_skills_line(doc, "Engine Development", "C++20, Vulkan, ECS architecture, visual scripting, shader authoring")
    add_skills_line(doc, "Gameplay Systems", "Encounters, interactive objects, narrative scripting, state machines, behavior trees, combat systems, 3D math")
    add_skills_line(doc, "Web Development", "Next.js, React, HTML/CSS, SEO")
    add_skills_line(doc, "Creative Coding", "p5.js, Processing")
    add_skills_line(doc, "Design Tools", "Blender, Photoshop")
    add_skills_line(doc, "Audio Integration", "FMOD, Pro Tools, Logic, Ableton")

    add_section_heading(doc, "Experience")

    add_role(doc, "NYU Game Center (Tisch) - Code Help Desk", "2025 - Present")
    add_bullet(doc, "Help BFA and MFA students debug and fix projects across Unity, Unreal, Processing, and GameMaker")
    add_bullet(doc, "Troubleshoot broken workflows and translate fixes into clear next steps")

    add_role(doc, "NYU Game Center - Studio Game Director & Producer", "Spring 2025")
    add_bullet(doc, "Led day-to-day production for a multidisciplinary game team from prototype through milestone builds")
    add_bullet(doc, "Coordinated playtests, translated findings into prioritized next steps, kept scope aligned with deadlines")
    add_bullet(doc, "Managed build stability, presentation flow, and last-mile polish for demo readiness")

    add_role(doc, "Custom Construction Website - Designer & Programmer", "2024")
    add_bullet(doc, "Built a company website with editable content tools and improved SEO")

    add_section_heading(doc, "Projects")

    add_role(doc, "TEGE - The Enjin Game Engine", "Solo Developer | C++20, Vulkan")
    add_bullet(doc, "Architected and shipped a 100k+ line aesthetic-first game engine with a complete ImGui editor, 70+ ECS components, and build pipeline")
    add_bullet(doc, "Implemented a Vulkan PBR renderer with cascaded shadow maps, clustered forward lighting, GPU occlusion culling, and full hardware ray tracing")
    add_bullet(doc, "Integrated dual physics backends (Jolt 5.2.0, Box2D 3.0.0), AngelScript scripting (721 bindings), and blueprint-style visual scripting with 146+ nodes")
    add_bullet(doc, "Built a 30+ effect post-processing stack, retro rendering modes (PSX, Dreamcast), and 44 starter templates")
    add_bullet(doc, "Developed accessibility suite: 8 colorblind corrections, remappable input, screen reader, switch access, dwell-click, dyslexia mode")
    add_bullet(doc, "Published under BSL 1.1 with full documentation, installer, and standalone game player")

    add_role(doc, "RhodesTrip", "Producer / Designer / Programmer / Audio | Unity, FMOD, Ink, Blender")
    add_bullet(doc, "Built narrative state machine, cameras, vehicle controls for a third-person driving + exploration game")
    add_bullet(doc, "Owned prototype slice, iteration based on playtests, cross-discipline coordination (team of 4)")

    add_role(doc, "IRIS (In-Progress)", "Solo Developer | Unity")
    add_bullet(doc, "Third-person environmental exploration with physics-based flower trimming and arranging")
    add_bullet(doc, "Prototyped tactile interactions; iterated on feel and reliability")

    add_role(doc, "Blackout", "Solo Developer | Unity, FMOD")
    add_bullet(doc, "Sightless 2D platformer: audio-driven feedback, navigation cues, mix logic, accessibility hooks, debug tools")
    add_bullet(doc, "Published research whitepaper on designing for blind play")

    add_role(doc, "Cold Flame", "Solo Developer | Unity")
    add_bullet(doc, "Autobiographical physics game; 9-minute experience using physics mechanics as emotional metaphor")

    add_role(doc, "Purrfect Pizza VR", "Design / Music / Audio Programming (Team of 6) | Unity, VR Toolkit, Logic")
    add_bullet(doc, "VR game gamifying formal language constructs through pizza-making for cat bosses")

    add_role(doc, "Prototype Studio Gallery", "11 Weekly Prototypes | Unity")
    add_bullet(doc, "One per week; text adventures, physics toys, 3D environments, naval strategy, abstract interactions")

    add_education(doc)

    doc.save(os.path.join(OUT_DIR, "marty-scott-programming.docx"))
    print("Built: marty-scott-programming.docx")


# ============================================================
# 3. GAME DEVELOPMENT
# ============================================================
def build_game_dev():
    doc = make_doc()
    add_header(doc)
    add_summary(doc,
        "A game designer and programmer who works where narrative meets mechanics and "
        "aesthetics meet systems. I prototype early, iterate on feel, and build end-to-end - "
        "from encounter logic and narrative state machines to camera systems and audio-driven "
        "accessibility. I also built an aesthetic-first game engine in C++/Vulkan. "
        "15+ years of audio makes me sensitive to timing and what a moment reads like."
    )

    add_section_heading(doc, "Skills")
    add_skills_line(doc, "Game Design", "Greyboxing, rapid prototyping, play-testing, iterative design, system balancing, tuning")
    add_skills_line(doc, "Gameplay Systems", "Encounters, interactive objects, narrative scripting, state machines, behavior trees, combat systems, 3D math")
    add_skills_line(doc, "Languages", "C#, C++, Python, JavaScript, Java, SQL")
    add_skills_line(doc, "Engines / Tools", "Unity, Unreal, GameMaker, TEGE (open source, BSL 1.1)")
    add_skills_line(doc, "Engine Development", "C++20, Vulkan, ECS architecture, visual scripting, shader authoring")
    add_skills_line(doc, "Audio", "FMOD, Pro Tools, Logic, Ableton (15+ years)")

    add_section_heading(doc, "Experience")

    add_role(doc, "NYU Game Center - Studio Game Director & Producer", "Spring 2025")
    add_bullet(doc, "Led production for a multidisciplinary game team, owning schedules and task tracking from prototype through milestone builds")
    add_bullet(doc, "Coordinated playtests and feedback capture, translated findings into prioritized next steps")
    add_bullet(doc, "Served as communication hub across design, engineering, art, and audio")
    add_bullet(doc, "Managed demo readiness: build stability, presentation flow, last-mile polish")

    add_role(doc, "NYU Game Center (Tisch) - Code Help Desk", "2025 - Present")
    add_bullet(doc, "Help BFA and MFA students with code issues across Unity, Unreal, Processing, and GameMaker")
    add_bullet(doc, "Troubleshoot broken projects and walk students through fixes")

    add_role(doc, "NYU Game Center Conference - Speaker", "2025")
    add_bullet(doc, "Presented on game design pedagogy and accessibility")

    add_section_heading(doc, "Projects")

    add_role(doc, "RhodesTrip", "Producer / Designer / Programmer / Audio | Unity, FMOD, Ink, Blender")
    add_bullet(doc, "Third-person driving + exploration with conversation and puzzle elements (team of 4)")
    add_bullet(doc, "Built narrative state machine, cameras, vehicle controls")
    add_bullet(doc, "Designed encounters, composed music, coordinated cross-discipline team")
    add_bullet(doc, "Owned prototype slice, iterated based on playtests")

    add_role(doc, "IRIS (In-Progress)", "Solo Developer | Unity")
    add_bullet(doc, "Third-person environmental exploration with physics-based flower trimming and arranging")
    add_bullet(doc, "Prototyped tactile interactions; iterated on feel and reliability")

    add_role(doc, "Blackout", "Solo Developer | Unity, FMOD")
    add_bullet(doc, "Sightless 2D platformer built around auditory navigation with no visual cues")
    add_bullet(doc, "Designed audio-driven feedback and progression: navigation cues, mix logic, accessibility hooks")
    add_bullet(doc, "Published research whitepaper on designing for blind play")

    add_role(doc, "Cold Flame", "Solo Developer | Unity")
    add_bullet(doc, "Autobiographical physics game about life, loss, and chopping wood")
    add_bullet(doc, "9-minute experience with voiceover, using physics mechanics as emotional metaphor")

    add_role(doc, "TEGE - The Enjin Game Engine", "Solo Developer | C++20, Vulkan")
    add_bullet(doc, "Designed, directed, and shipped an aesthetic-first open source game engine with a complete editor, 70+ ECS components, dual physics backends (Jolt/Box2D), and AngelScript scripting with 721 bindings")
    add_bullet(doc, "Full ray tracing pipeline, 30+ post-processing effects, retro rendering modes, and 44 starter templates")
    add_bullet(doc, "Accessibility suite with 8 colorblind modes, screen reader support, switch access, and motor accessibility")

    add_role(doc, "Purrfect Pizza VR", "Design / Music / Audio Programming (Team of 6) | Unity, VR Toolkit, Logic")
    add_bullet(doc, "VR game gamifying formal language constructs through pizza-making for cat bosses")

    add_role(doc, "Prototype Studio Gallery", "11 Weekly Prototypes | Unity")
    add_bullet(doc, "One per week; text adventures, physics toys, 3D environments, naval strategy, emotional vignettes")

    add_education(doc)

    doc.save(os.path.join(OUT_DIR, "marty-scott-game-development.docx"))
    print("Built: marty-scott-game-development.docx")


# ============================================================
# 4. COMBINED
# ============================================================
def build_combined():
    doc = make_doc()
    add_header(doc)
    add_summary(doc,
        "A game designer, programmer, and educator with a CS degree and an MFA in Game Design "
        "from NYU. I build playable prototypes where narrative and mechanics meet, write "
        "aesthetic-first game engines, design for accessibility from the start, and teach code "
        "as creative material. 15+ years of audio across music, podcasts, and game audio."
    )

    add_section_heading(doc, "Skills")
    add_skills_line(doc, "Languages", "C#, C++, Java, Python, JavaScript, SQL")
    add_skills_line(doc, "Engines / Tools", "Unity, Unreal, GameMaker, TEGE (open source, BSL 1.1)")
    add_skills_line(doc, "Engine Development", "C++20, Vulkan, ECS architecture, visual scripting, shader authoring")
    add_skills_line(doc, "Game Design", "Greyboxing, rapid prototyping, play-testing, iterative design, system balancing")
    add_skills_line(doc, "Gameplay Systems", "Encounters, interactive objects, narrative scripting, state machines, behavior trees, 3D math")
    add_skills_line(doc, "Pedagogy", "Curriculum design, scaffolded instruction, studio-based teaching, accessibility-centered design")
    add_skills_line(doc, "Audio", "FMOD, Pro Tools, Logic, Ableton (15+ years)")
    add_skills_line(doc, "Video / Design", "Premiere Pro, Photoshop - editing, thumbnails, social assets")
    add_skills_line(doc, "Operations", "Scheduling, outreach, event logistics, documentation, Discord community management")

    add_section_heading(doc, "Experience")

    add_role(doc, "NYU Game Center (Tisch) - Code Help Desk", "2025 - Present")
    add_bullet(doc, "Help BFA and MFA students with code issues across Unity, Unreal, Processing, and GameMaker")
    add_bullet(doc, "Troubleshoot broken projects and walk students through fixes")

    add_role(doc, "NYU Game Center - Studio Game Director & Producer", "Spring 2025")
    add_bullet(doc, "Led production for a multidisciplinary game team from prototype through milestone builds")
    add_bullet(doc, "Coordinated playtests, managed build stability, and kept scope aligned with deadlines")

    add_role(doc, "NYU Game Center Conference - Speaker", "2025")
    add_bullet(doc, "Presented on game design pedagogy and accessibility")

    add_role(doc, "Fib Likley (YouTube) - Creator & Editor", "2019 - Present")
    add_bullet(doc, "Gameplay-driven videos with clear hooks, clean structure, and tight pacing")

    add_role(doc, "The Palisades - Recording Studio Manager & Operations", "2013 - 2019")
    add_bullet(doc, "Coordinated scheduling for 50+ bands, managed 200+ podcast sessions")
    add_bullet(doc, "Produced monthly public events, managed on-site logistics and vendor coordination")

    add_role(doc, "Music & Arts - Music Teacher", "2012 - 2018")
    add_bullet(doc, "Taught piano, guitar, and sound production across age groups and skill levels")

    add_section_heading(doc, "Projects")

    add_role(doc, "TEGE - The Enjin Game Engine", "Solo Developer | C++20, Vulkan")
    add_bullet(doc, "Solo-developed an aesthetic-first open source game engine featuring 70+ ECS components, a full ImGui editor, ray tracing pipeline, and comprehensive accessibility suite. BSL 1.1")

    add_role(doc, "RhodesTrip", "Producer / Designer / Programmer / Audio | Unity, FMOD, Ink, Blender")
    add_bullet(doc, "Third-person driving + exploration; narrative state machine, cameras, vehicle controls (team of 4)")

    add_role(doc, "Blackout", "Solo Developer | Unity, FMOD")
    add_bullet(doc, "Sightless 2D platformer with audio-driven navigation; published research whitepaper")

    add_role(doc, "IRIS (In-Progress)", "Solo Developer | Unity")
    add_bullet(doc, "Physics-based environmental exploration; prototyping tactile interactions")

    add_role(doc, "Cold Flame", "Solo Developer | Unity")
    add_bullet(doc, "Autobiographical physics game; code as personal expression")

    add_section_heading(doc, "Pedagogical Work")
    add_role(doc, "VG101 Knowledge Base", "marty64.net/vg101 | In Development")
    add_bullet(doc, "Open text for teaching videogame design as its own medium (Gesture + Aesthetic Heritage framework)")
    add_bullet(doc, "Teaching cycle: Play - Name - Make - Reflect; three-register system for layered reading")

    add_education(doc)

    doc.save(os.path.join(OUT_DIR, "marty-scott-combined.docx"))
    print("Built: marty-scott-combined.docx")


# ============================================================
if __name__ == "__main__":
    build_teaching()
    build_programming()
    build_game_dev()
    build_combined()
    print("\nAll DOCX resumes generated.")
